import os
from googlesearch import search
from datetime import datetime

from flask import Flask, send_file, request
from docx import Document


app = Flask(__name__)


def get_google_search_results(query, num_results=10):
    search_results = []
    try:
        # Perform the Google search
        search_results = list(search(query, num_results=num_results))
    except Exception as e:
        print("An error occurred:", e)
    
    return search_results


@app.route('/s/<query>')
def search_route(query):
    num_results = request.args.get('n')
    results = get_google_search_results(query, num_results=int(num_results) if num_results else 10)
    doc = Document()
    doc.add_heading('Search results for query: ' + query, 0)
    for idx, result in enumerate(results, start=1):
        doc.add_paragraph(f'{idx}. {result}')

    current_datetime = datetime.now()
    current_datetime_string = current_datetime.strftime('%Y-%m-%d_%H:%M:%S')
    filename=f'{current_datetime_string}_{query.replace(" ", "")}.docx'
    doc.save(filename)
    return send_file(filename, as_attachment=True)

def main():
    app.run(port=int(os.environ.get('PORT', 8000)))

if __name__ == "__main__":
    main()
